import asyncio
import io

from app.services.parser.base import BaseFileParser


class DocxParser(BaseFileParser):

    async def parse(self, data: bytes) -> str:
        return await asyncio.to_thread(self._extract_text, data)

    @staticmethod
    def _extract_text(data: bytes) -> str:
        from docx import Document

        doc = Document(io.BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)
